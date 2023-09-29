##### `ghost.py`
##### Ghostwriter
##### Open-Source, hosted on https://github.com/DrBenjamin/Ghostwriter
##### Please reach out to ben@benbox.org for any questions
#### Loading needed Python libraries
import streamlit as st
from streamlit.connections import SQLConnection
import io
import pandas as pd
#import pymysql
#pymysql.install_as_MySQLdb()
#from sqlalchemy.sql import text
import openai
import deepl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH




#### Functions
### Function: export_docx = Pandas dataframe to MS Word file (docx)
def export_docx(text, docx_file_name = 'Ghostwriter.docx'):
    document = Document()

    # Adding header
    document.add_heading('Ghostwriter', level = 0)

    # Adding text
    document.add_heading('Test', level = 1)
    paragraph = document.add_paragraph()
    data = text['SENTENCE'].tolist()
    for d in data:
        paragraph.add_run(d + " ")
          
    # Create a Word file using python-docx as engine
    buffer = io.BytesIO()
    document.save(buffer)
    
    # Download Button
    st.toast('Your document is ready!', icon = '📃')
    st.download_button(label = 'Download Word document', data = buffer, file_name = docx_file_name, mime = "application/vnd.openxmlformats")



### Function: trans = DeepL translation
def trans(input, target_lang):
    translator = deepl.Translator(st.secrets["deepl"]["key"])
    result = translator.translate_text(input, target_lang = target_lang)
    return result



### Function: lastID = checks for last ID number in Table (to add data after)
def lastID(url):
    rows = connection.query("SELECT MAX(ID) FROM " + url + ";")

    # Check for ID
    try:
        id = len(rows) + 1
    except:
        id = 1

    # Return ID
    return id




#### Main
### Get data from the databank(s)
# Open databank connection
connection = st.experimental_connection(name = 'ghost', type = 'sql', autocommit = True)

# Get employee data
text = connection.query("SELECT ID, SENTENCE FROM `ghostwriter`.`document`;")
text = text.set_index('ID')
st.write(text)



### Input Form
with st.form(key = 'Ghostwriter'):
    st.header("Ghostwriter")
    st.write("Welcome to Ghostwriter, a tool to help you write better. Ghostwriter uses the power of AI to help you write better, faster, and more efficiently.")
    sentence = st.text_input("Please enter your text here:")

    answer = ''
    submitted = st.form_submit_button('Submit')
    if submitted:
        # Set API key
        openai.api_key = st.secrets['openai']['key']

        # OpenAI processing
        model = 'gpt-3.5-turbo' #model == 'gpt-4-0613', 'gpt-4' or 'gpt-3.5-turbo'
        response_answer = openai.ChatCompletion.create(
            model = model,
            messages = [{"role": "system", "content": 'Please rewrite the user sentence in a more academic form, but stick to the information given.'},
                        {"role": "user", "content": sentence},])
                                
        # Cleaning answer
        answer = response_answer['choices'][0]['message']['content'].lstrip()
        answer = answer.replace("'", "")
        answer = answer.replace('"', '')
        st.write(answer)
        st.write(trans(answer, "DE"))

# Export Word document
id = lastID("`ghostwriter`.`document`")
if answer == '':
    answer = 'Test ' + str(id)
with connection.session as session:
    session.execute('INSERT INTO document (ID, SENTENCE) VALUES(:idd, :sen);', params = dict(idd = id, sen = answer))
    session.commit()


try:
    export_docx(text)
    
except:
    print('No answer so far')
